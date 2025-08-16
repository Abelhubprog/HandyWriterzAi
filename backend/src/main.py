"""
Revolutionary FastAPI application for HandyWriterz backend.
Production-ready with comprehensive error handling and resilience patterns.
"""
from dotenv import load_dotenv

load_dotenv()

import asyncio
import json
import logging
import os
import time
import uuid
from contextlib import asynccontextmanager
from typing import Dict, Any, Optional, List

import redis.asyncio as redis
import uvicorn
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Depends, Request, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse, FileResponse
from fastapi.exceptions import RequestValidationError
from fastapi.staticfiles import StaticFiles
from langchain_core.messages import HumanMessage
import logging
import logging.config
from src.config import get_settings, setup_logging

# Setup structured logging using configuration system
settings = get_settings()
setup_logging(settings)
logger = logging.getLogger(__name__)

# Initialize multi-provider AI system
from src.models.factory import initialize_factory, get_factory, get_provider
from src.models.base import ModelRole

# Initialize AI provider factory with API keys from settings
ai_provider_factory = initialize_factory({
    "gemini": settings.gemini_api_key,
    "openai": settings.openai_api_key,
    "anthropic": settings.anthropic_api_key,
    "perplexity": settings.perplexity_api_key
})

logger.info("🤖 Multi-provider AI system initialized")

import asyncio
import json
import os
import time
import uuid
from contextlib import asynccontextmanager
from typing import Dict, Any, Optional, List

import redis.asyncio as redis
import uvicorn
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Depends, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse, FileResponse
from fastapi.exceptions import RequestValidationError
from fastapi.staticfiles import StaticFiles
from langchain_core.messages import HumanMessage
from pydantic import BaseModel

# Import agent system for HandyWriterz workflow
from src.agent.handywriterz_state import HandyWriterzState
from src.agent.base import UserParams

# Simple system removed - all requests use advanced HandyWriterz system
SIMPLE_SYSTEM_AVAILABLE = False
logger.info("✅ Simple system permanently disabled - using advanced system only")

# Import routing system from the canonical location only (no silent fallbacks)
from src.agent.routing.unified_processor import UnifiedProcessor

from src.db.database import (
    get_user_repository, get_conversation_repository,
    get_document_repository, db_manager
)

# Import admin routes
from src.routes.admin_models import router as admin_models_router
from src.routes.admin_credits import router as admin_credits_router
from src.routes.credits import router as user_credits_router
from src.routes.stream import router as stream_router
from src.routes.stream_extras import router as stream_extras_router
from src.routes.chat_gateway import chat_gateway_router
from src.routes.health import router as health_router
from src.api.autonomy_v2 import router as autonomy_v2_router
from src.api.files import router as files_router
from src.services.error_handler import (
    error_handler, ErrorContext, ErrorCategory, ErrorSeverity,
    with_error_handling, with_circuit_breaker, with_retry
)
from src.services.security_service import (
    security_service, get_current_user, require_authorization,
    require_rate_limit, validate_input
)

# Import new Phase 1 components
from src.models.registry import initialize_registry, get_registry
from src.services.budget import get_budget_guard
from src.services.logging_context import setup_correlation_logging
from src.services.feature_validator import get_feature_validator
from src.middleware.error_middleware import (
    error_middleware, global_exception_handler
)
from src.middleware.security_middleware import (
    security_middleware, csrf_middleware
)
from scripts import init_database

# Initialize model registry with strict mode in production
try:
    strict_mode = os.getenv("FEATURE_REGISTRY_ENFORCED", "false").lower() == "true"
    registry_config_path = "src/config/model_config.yaml"
    pricing_config_path = "src/config/price_table.json"

    initialize_registry(
        model_config_path=registry_config_path,
        price_table_path=pricing_config_path,
        strict=strict_mode
    )

    logger.info(f"✅ ModelRegistry initialized (strict_mode: {strict_mode})")

except Exception as e:
    if os.getenv("FEATURE_REGISTRY_ENFORCED", "false").lower() == "true":
        logger.error(f"❌ ModelRegistry required but failed to initialize: {e}")
        raise RuntimeError(f"Production requires valid ModelRegistry: {e}")
    else:
        logger.warning(f"⚠️  ModelRegistry initialization failed (non-strict mode): {e}")

# Async Redis client for SSE
# Normalize ALLOWED_ORIGINS and other env from .env.* files; keep defaults safe
redis_client = redis.from_url(os.getenv("REDIS_URL", "redis://localhost:6379"), decode_responses=True)


# Initialize unified processor
unified_processor = UnifiedProcessor(simple_available=False, advanced_available=True)


from src.db.models import Base
from sqlalchemy import create_engine

# Create database tables on startup, outside of the lifespan manager
# to ensure they are created even if the lifespan event doesn't fire
# in all execution contexts (like some test runners or script executions).
try:
    db_url = os.getenv("DATABASE_URL")
    if db_url:
        if db_url.startswith("postgres://"):
            db_url = db_url.replace("postgres://", "postgresql://", 1)
        engine = create_engine(db_url)
        Base.metadata.create_all(bind=engine)
        logger.info("✅ Database tables created successfully on startup.")
    else:
        logger.warning("⚠️ DATABASE_URL not set, skipping table creation on startup.")
except Exception as e:
    logger.error(f"❌ Failed to create database tables on startup: {e}")


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Manage application lifespan with comprehensive health checks."""
    logger.info("Starting HandyWriterz Revolutionary Backend...")

    # store start time
    app.state.start_time = time.time()

    # Initialize database with system prompts (skip if using SQLite fallback or missing DB)
    try:
        db_url = os.getenv("DATABASE_URL", "")
        if db_url and not db_url.startswith("sqlite"):
            await init_database.main()
            logger.info("✅ Database initialized with system prompts.")
        else:
            logger.info("ℹ️ Skipping init_database for SQLite/no DB in dev.")
    except Exception as e:
        logger.error(f"❌ Failed to initialize database with system prompts: {e}")

    # Initialize systems with error handling
    startup_errors = []

    # Setup correlation logging
    try:
        setup_correlation_logging()
        logger.info("✅ Correlation logging initialized")
    except Exception as e:
        startup_errors.append(f"Correlation logging setup failed: {e}")
        logger.error(f"❌ Correlation logging setup failed: {e}")

    # Initialize model registry
    try:
        model_config_path = "src/config/model_config.yaml"
        price_table_path = "src/config/price_table.json"
        strict_mode = os.getenv("FEATURE_REGISTRY_ENFORCED", "false").lower() == "true"

        registry = initialize_registry(model_config_path, price_table_path, strict_mode)
        if registry.validate():
            logger.info("✅ Model registry initialized and validated")
        else:
            startup_errors.append("Model registry validation failed")
            logger.error("❌ Model registry validation failed")
    except Exception as e:
        startup_errors.append(f"Model registry initialization failed: {e}")
        logger.error(f"❌ Model registry initialization failed: {e}")

    # Initialize budget guard
    try:
        budget_guard = get_budget_guard()
        logger.info("✅ Budget guard initialized")
    except Exception as e:
        startup_errors.append(f"Budget guard initialization failed: {e}")
        logger.error(f"❌ Budget guard initialization failed: {e}")

    # Diagnostic: pgvector + SQLAlchemy support check (common Py3.12 pitfall)
    try:
        import sys as _sys
        pyver = _sys.version.split()[0]
        try:
            from pgvector import sqlalchemy as _pgvec_sa  # type: ignore
            logger.info(f"✅ pgvector.sqlalchemy available (Python {pyver})")
        except Exception as _e:
            logger.warning(
                "⚠️  pgvector.sqlalchemy not available. If you need vector DB features, use Python 3.11 and pgvector>=0.4.5, "
                "or add a fallback TypeDecorator. Current Python: %s; error: %s", pyver, _e
            )
    except Exception:
        # Don't block startup on diagnostics
        pass

    # Test Redis connection (fail fast in managed deployments like Railway)
    try:
        # Use a timeout via wait_for to avoid long hangs
        await asyncio.wait_for(redis_client.ping(), timeout=float(os.getenv("REDIS_PING_TIMEOUT", "2")))
        logger.info("✅ Redis connection successful")
    except Exception as e:
        logger.error(f"❌ Redis connection failed: {e}")
        raise RuntimeError(f"Redis connection failed: {e}")

    # Warm up SSE Service (lazy client) to ensure Redis availability for publishers
    try:
        from src.services.sse_service import get_sse_service
        sse_service = get_sse_service()
        ok = await sse_service.ping()
        if ok:
            logger.info("✅ SSEService ping successful")
        else:
            startup_errors.append("SSEService ping failed")
            logger.warning("⚠️ SSEService ping failed")
    except Exception as e:
        startup_errors.append(f"SSEService init failed: {e}")
        logger.error(f"❌ SSEService init failed: {e}")

    # Test Database connection
    try:
        if db_manager.health_check():
            logger.info("✅ Database connection successful")
        else:
            startup_errors.append("Database health check failed")
            logger.error("❌ Database health check failed")
    except Exception as e:
        startup_errors.append(f"Database connection failed: {e}")
        logger.error(f"❌ Database connection failed: {e}")

    # Test Error Handler
    try:
        await error_handler.get_error_statistics()
        logger.info("✅ Error handler initialized successfully")
    except Exception as e:
        startup_errors.append(f"Error handler initialization failed: {e}")
        logger.error(f"❌ Error handler initialization failed: {e}")

    # Log startup summary
    if startup_errors:
        logger.warning(f"Backend started with {len(startup_errors)} issues:")
        for error in startup_errors:
            logger.warning(f"  - {error}")
    else:
        logger.info("🚀 All systems operational - HandyWriterz backend ready!")

    try:
        yield
    finally:
        logger.info("Shutting down HandyWriterz backend...")

        # Close database connections
        try:
            db_manager.close()
            logger.info("✅ Database connections closed")
        except Exception as e:
            logger.error(f"❌ Error closing database: {e}")

        # Close Redis connections
        try:
            await redis_client.close()
            logger.info("✅ Redis connections closed")
        except Exception as e:
            logger.error(f"❌ Error closing Redis: {e}")

        # Close SSEService client if initialized
        try:
            from src.services.sse_service import get_sse_service
            sse_service = get_sse_service()
            await sse_service.close()
            logger.info("✅ SSEService closed")
        except Exception as e:
            logger.warning(f"⚠️ SSEService close encountered an issue: {e}")

        logger.info("🔄 HandyWriterz backend shutdown complete")


# Create FastAPI app with enhanced configuration
app = FastAPI(
    title="Unified AI Platform - Revolutionary API",
    description="🚀 Intelligent Multi-Agent System: Simple Gemini + Advanced HandyWriterz with Automatic Routing",
    version="2.0.0-unified",
    lifespan=lifespan,
    docs_url="/docs",
    redoc_url="/redoc",
    openapi_url="/openapi.json"
)

# Add Security Middleware (first layer of protection)
app.add_middleware(security_middleware)

# Add CSRF Protection Middleware
app.add_middleware(csrf_middleware)

# Add Error Handling Middleware
app.add_middleware(error_middleware)

# Add CORS middleware (parse ALLOWED_ORIGINS from env; supports comma/semicolon/json)
def _parse_allowed_origins() -> list[str]:
    raw = os.getenv("ALLOWED_ORIGINS", "")
    if not raw:
        return [
            "http://localhost:3000",
            "http://localhost:3001",
            "http://localhost:5173",
            "https://handywriterz.vercel.app",
        ]
    raw = raw.strip()
    # Try JSON array first
    if raw.startswith("["):
        try:
            arr = json.loads(raw)
            return [str(x).strip() for x in arr if str(x).strip()]
        except Exception:
            pass
    # Fallback to split by comma or semicolon
    parts = []
    for sep in [",", ";"]:
        if sep in raw:
            parts = [p.strip() for p in raw.split(sep)]
            break
    if not parts:
        parts = [raw.strip()]
    return [p for p in parts if p]

app.add_middleware(
    CORSMiddleware,
    allow_origins=_parse_allowed_origins(),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["X-Request-ID", "X-Response-Time", "X-Error-ID", "X-Security-Middleware"]
)

# Add global exception handlers
@app.exception_handler(RequestValidationError)
async def validation_exception_handler(req: Request, exc: RequestValidationError):
    return JSONResponse(
        status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
        content={"detail": exc.errors(), "body": exc.body},
    )

app.add_exception_handler(HTTPException, global_exception_handler.handle_http_exception)

# Include admin routes
app.include_router(admin_models_router)
app.include_router(admin_credits_router)
app.include_router(user_credits_router)
app.include_router(admin_credits_router)
app.include_router(files_router, prefix="/api")
from src.api.billing import router as billing_router
from src.api.profile import router as profile_router
from src.api.usage import router as usage_router
from src.api.memory import router as memory_router

app.include_router(billing_router, prefix="/api")
app.include_router(profile_router, prefix="/api")
app.include_router(usage_router, prefix="/api")
app.include_router(memory_router, prefix="/api", tags=["memory"])  # Expose memory endpoints
app.include_router(health_router)


# Include payment system routes
# Payments and payout routes are guarded due to model mismatches with WalletEscrow schema.
# Enable via FEATURE_PAYMENTS_ENABLED=true when models are aligned.
from src.api.checker import router as checker_router
app.include_router(checker_router)

if os.getenv("FEATURE_PAYMENTS_ENABLED", "false").lower() == "true":
    try:
        from src.api.payments import router as payments_router
        from src.api.payout import router as payout_router
        app.include_router(payments_router)
        app.include_router(payout_router)
        logger.info("✅ Payments routes enabled by FEATURE_PAYMENTS_ENABLED")
    except Exception as e:
        logger.warning(f"⚠️ Payments routes disabled due to import/runtime error: {e}")

# Include Workbench routes
from src.api.workbench import router as workbench_router
from src.api.workbench_auth import router as workbench_auth_router
from src.api.workbench_ingestion import router as workbench_ingestion_router
# from src.api.workbench_admin import router as workbench_admin_router # Temporarily commented out until implemented

app.include_router(workbench_router, prefix="/api", tags=["workbench"])
app.include_router(workbench_auth_router, prefix="/api", tags=["workbench-auth"])
app.include_router(workbench_ingestion_router, prefix="/api", tags=["workbench-ingestion"])
# app.include_router(workbench_admin_router, prefix="/api", tags=["workbench-admin"]) # Temporarily commented out

# SSE unified stream router
app.include_router(stream_router)
app.include_router(stream_extras_router)
app.include_router(chat_gateway_router)

# Feature-flagged mount for Autonomy V2 router under /api/v2
try:
    if settings.enable_autonomy_v2:
        app.include_router(autonomy_v2_router, prefix="/api")
        logger.info("✅ Autonomy V2 router mounted under /api/v2")
    else:
        logger.info("ℹ️ Autonomy V2 disabled (ENABLE_AUTONOMY_V2=false)")
except Exception as e:
    logger.warning(f"⚠️ Failed to mount Autonomy V2 router: {e}")


# Mount static files for serving the SvelteKit frontend
static_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "static")
if os.path.exists(static_dir):
    app.mount("/static", StaticFiles(directory=static_dir), name="static")
    logger.info(f"✅ Static files mounted from {static_dir}")
else:
    logger.warning(f"⚠️ Static directory not found: {static_dir}")

# Mount pyodide files if they exist
pyodide_dir = os.path.join(static_dir, "pyodide") if os.path.exists(static_dir) else None
if pyodide_dir and os.path.exists(pyodide_dir):
    app.mount("/pyodide", StaticFiles(directory=pyodide_dir), name="pyodide")
    logger.info(f"✅ Pyodide files mounted from {pyodide_dir}")

# Mount SvelteKit build directory for frontend serving
build_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "build")
if os.path.exists(build_dir):
    app.mount("/app", StaticFiles(directory=build_dir), name="frontend_build")
    logger.info(f"✅ SvelteKit build mounted from {build_dir}")
else:
    logger.warning(f"⚠️ SvelteKit build directory not found: {build_dir}")


# Pydantic models
class WritingRequest(BaseModel):
    """Request model for academic writing."""
    prompt: str
    user_params: Dict[str, Any]
    auth_token: Optional[str] = None
    payment_transaction_id: Optional[str] = None
    uploaded_file_urls: List[str] = []


class WritingResponse(BaseModel):
    """Response model for writing request."""
    conversation_id: str
    status: str
    message: str


class HealthResponse(BaseModel):
    """Health check response."""
    status: str
    timestamp: float
    version: str


# Health check endpoints
@app.get("/health", response_model=HealthResponse)
@with_error_handling(ErrorCategory.SYSTEM, ErrorSeverity.MEDIUM)
async def health_check():
    """Comprehensive health check endpoint."""
    return HealthResponse(
        status="healthy",
        timestamp=time.time(),
        version="2.0.0-unified"
    )


# Enhanced system status endpoint for unified platform
@app.get("/api/status")
@with_error_handling(ErrorCategory.SYSTEM, ErrorSeverity.LOW)
async def unified_system_status():
    """
    🎯 Unified System Status - Enhanced with routing information
    Shows the status of both simple and advanced systems with routing capabilities.
    """
    try:
        # Get routing statistics
        routing_stats = unified_processor.router.get_routing_stats()

        # Test system availability
        simple_status = "permanently_disabled"
        advanced_status = "available"  # HandyWriterz is always available in this context

        # Get Redis status
        redis_status = "unknown"
        try:
            await redis_client.ping()
            redis_status = "healthy"
        except:
            redis_status = "unhealthy"

        # Feature flags snapshot
        _settings = get_settings() if get_settings else None
        feature_flags = {
            "sse_publisher_unified": os.getenv("FEATURE_SSE_PUBLISHER_UNIFIED", "false").lower() == "true",
            "params_normalization": os.getenv("FEATURE_PARAMS_NORMALIZATION", "false").lower() == "true",
            "double_publish_sse": os.getenv("FEATURE_DOUBLE_PUBLISH_SSE", "false").lower() == "true",
            "registry_enforced": os.getenv("FEATURE_REGISTRY_ENFORCED", "false").lower() == "true",
            "search_adapter": os.getenv("FEATURE_SEARCH_ADAPTER", "true").lower() == "true",
            "turnitin_hitl_enabled": os.getenv("FEATURE_TURNITIN_HITL_ENABLED", "false").lower() == "true", # New flag
        }

        # Override with settings if available
        if _settings:
            feature_flags.update({
                "sse_publisher_unified": getattr(_settings, "feature_sse_publisher_unified", feature_flags["sse_publisher_unified"]),
                "params_normalization": getattr(_settings, "feature_params_normalization", feature_flags["params_normalization"]),
                "double_publish_sse": getattr(_settings, "feature_double_publish_sse", feature_flags["double_publish_sse"]),
                "registry_enforced": getattr(_settings, "feature_registry_enforced", feature_flags["registry_enforced"]),
                "search_adapter": getattr(_settings, "feature_search_adapter", feature_flags["search_adapter"]),
                "turnitin_hitl_enabled": getattr(_settings, "feature_turnitin_hitl_enabled", feature_flags["turnitin_hitl_enabled"]),
            })

        # Get model registry status
        registry_status = "unknown"
        registry_models = 0
        try:
            registry = get_registry()
            if registry._loaded:
                registry_status = "loaded"
                registry_models = len(registry.get_all_models())
            else:
                registry_status = "not_loaded"
        except:
            registry_status = "error"

        # Get budget guard status
        budget_status = "unknown"
        try:
            budget_guard = get_budget_guard()
            budget_status = "active"
        except:
            budget_status = "error"

        # Get database status
        db_status = "unknown"
        try:
            if db_manager.health_check():
                db_status = "healthy"
            else:
                db_status = "unhealthy"
        except:
            db_status = "unavailable"

        return {
            "status": "operational",
            "version": "2.0.0-unified",
            "timestamp": time.time(),
            "platform": "Unified AI Platform",

            # System availability
            "systems": {
                "simple_gemini": {
                    "status": simple_status,
                    "description": "Quick responses for simple queries",
                    "capabilities": ["chat", "basic_research", "quick_answers"]
                },
                "advanced_handywriterz": {
                    "status": advanced_status,
                    "description": "Full academic writing workflow with 30+ agents",
                    "capabilities": [
                        "academic_writing", "research_swarms", "qa_swarms",
                        "writing_swarms", "citation_management", "turnitin_checking",
                        "learning_outcomes", "multi_agent_orchestration"
                    ]
                }
            },

            # Intelligent routing
            "routing": {
                "enabled": True,
                "algorithm": "complexity_analysis",
                "thresholds": routing_stats["thresholds"],
                "modes": routing_stats["routing_modes"],
                "capabilities": routing_stats["capabilities"]
            },

            # Infrastructure
            "infrastructure": {
                "redis": {
                    "status": redis_status,
                    "purpose": "Caching and SSE"
                },
                "database": {
                    "status": db_status,
                    "purpose": "User data and conversations"
                },
                "vector_storage": {
                    "status": "available",
                    "purpose": "Semantic search and embeddings"
                },
                "model_registry": {
                    "status": registry_status,
                    "models_loaded": registry_models,
                    "purpose": "Model ID mapping and pricing"
                },
                "budget_guard": {
                    "status": budget_status,
                    "purpose": "Cost control and abuse prevention"
                }
            },

            # Features
            "features": {
                "intelligent_routing": True,
                "multimodal_processing": True,
                "file_upload": True,
                "streaming_responses": True,
                "swarm_intelligence": advanced_status == "available",
                "academic_writing": advanced_status == "available",
                "web3_authentication": True,
                "real_time_collaboration": True,
                "citation_management": True,
                "plagiarism_checking": True,
                "flags": feature_flags
            },

            # API endpoints
            "endpoints": {
                "unified_chat": "/api/chat",
                "simple_chat": "removed - use /api/chat",
                "advanced_chat": "/api/chat/advanced",
                "academic_writing": "/api/write",
                "file_upload": "/api/upload",
                "streaming": "/api/stream/{conversation_id}",
                "documentation": "/docs",
                "workbench": "/api/workbench", # New endpoint
                "feature_status": "/api/features/status"
            },

            # Performance metrics
            "performance": {
                "routing_overhead": "< 50ms",
                "simple_response_time": "1-3 seconds",
                "advanced_response_time": "30-300 seconds",
                "hybrid_response_time": "30-300 seconds (parallel processing)"
            }
        }

    except Exception as e:
        logger.error(f"System status check failed: {e}")
        return {
            "status": "degraded",
            "version": "2.0.0-unified",
            "timestamp": time.time(),
            "error": str(e),
            "systems": {
                "simple_gemini": {"status": "unknown"},
                "advanced_handywriterz": {"status": "unknown"}
            }
        }


# Feature Status Endpoint for Disabled Services
@app.get("/api/features/status")
@with_error_handling(ErrorCategory.SYSTEM, ErrorSeverity.LOW)
async def get_features_status():
    """
    🔧 Feature Status Endpoint
    Returns structured information about disabled/enabled features with reasons and alternatives.
    """
    try:
        validator = get_feature_validator()
        return {
            "status": "success",
            "timestamp": time.time(),
            "features": validator.get_all_features_status()
        }
    except Exception as e:
        logger.error(f"Features status check failed: {e}")
        return {
            "status": "error",
            "timestamp": time.time(),
            "error": str(e),
            "features": {}
        }


# Admin Metrics Endpoint for Usage Tracking
@app.get("/api/admin/metrics")
@with_error_handling(ErrorCategory.SYSTEM, ErrorSeverity.LOW)
async def get_admin_metrics(
    current_user: Dict[str, Any] = Depends(require_authorization("admin")),
    tenant: Optional[str] = None,
    include_models: bool = True,
    include_providers: bool = True,
    time_window: str = "24h"
):
    """
    📊 Admin Metrics Endpoint
    Returns comprehensive usage metrics for administrators including costs, tokens, and performance.
    """
    try:
        # Get budget guard with Redis metrics
        budget_guard = get_budget_guard()

        # Base metrics structure
        metrics = {
            "timestamp": time.time(),
            "time_window": time_window,
            "tenant": tenant or "all",
            "summary": {},
            "usage": {},
            "costs": {},
            "performance": {},
            "providers": {},
            "models": {}
        }

        # Get usage summary from Redis-backed budget guard
        if tenant:
            usage_summary = budget_guard.get_usage_summary(tenant)
            metrics["usage"] = usage_summary
            metrics["summary"] = {
                "total_requests": usage_summary.get("total_requests", 0),
                "total_tokens": usage_summary.get("total_tokens", 0),
                "total_cost": usage_summary.get("total_cost", 0.0),
                "daily_spent": usage_summary.get("daily_spent", 0.0),
                "hourly_spent": usage_summary.get("hourly_spent", 0.0),
                "monthly_spent": usage_summary.get("monthly_spent", 0.0),
                "data_source": usage_summary.get("data_source", "unknown")
            }
        else:
            # Aggregate metrics for all tenants (would need Redis SCAN in production)
            metrics["summary"] = {
                "note": "Tenant-specific metrics available with ?tenant=<tenant_id>",
                "total_requests": "aggregation_not_implemented",
                "total_tokens": "aggregation_not_implemented",
                "total_cost": "aggregation_not_implemented"
            }

        # Provider performance metrics
        if include_providers:
            try:
                factory = get_factory()
                providers = factory.list_providers()

                for provider_name in providers:
                    provider = factory.get_provider(provider_name)
                    if provider:
                        metrics["providers"][provider_name] = {
                            "status": "available",
                            "default_model": getattr(provider, 'get_default_model', lambda: 'unknown')(),
                            "provider_type": provider.__class__.__name__
                        }

            except Exception as e:
                logger.warning(f"Failed to get provider metrics: {e}")
                metrics["providers"] = {"error": str(e)}

        # Model registry metrics
        if include_models:
            try:
                registry = get_registry()
                if registry._loaded:
                    all_models = registry.get_all_models()
                    metrics["models"] = {
                        "total_models": len(all_models),
                        "models_by_provider": {},
                        "registry_status": "loaded"
                    }

                    # Group models by provider
                    for model_id, model_info in all_models.items():
                        provider = model_info.provider
                        if provider not in metrics["models"]["models_by_provider"]:
                            metrics["models"]["models_by_provider"][provider] = []
                        metrics["models"]["models_by_provider"][provider].append({
                            "logical_id": model_id,
                            "actual_model": model_info.model_id,
                            "has_pricing": model_info.pricing is not None
                        })
                else:
                    metrics["models"] = {"registry_status": "not_loaded"}

            except Exception as e:
                logger.warning(f"Failed to get model metrics: {e}")
                metrics["models"] = {"error": str(e)}

        # System performance metrics
        metrics["performance"] = {
            "simple_system_available": False,  # Removed
            "advanced_system_available": True,
            "redis_status": "unknown",
            "database_status": "unknown"
        }

        # Test Redis connection
        try:
            await redis_client.ping()
            metrics["performance"]["redis_status"] = "healthy"
        except Exception:
            metrics["performance"]["redis_status"] = "unhealthy"

        # Test database connection
        try:
            if db_manager.health_check():
                metrics["performance"]["database_status"] = "healthy"
            else:
                metrics["performance"]["database_status"] = "unhealthy"
        except Exception:
            metrics["performance"]["database_status"] = "error"

        return metrics

    except Exception as e:
        logger.error(f"Admin metrics failed: {e}")
        raise HTTPException(
            status_code=500,
            detail={
                "error": "metrics_collection_failed",
                "message": str(e),
                "timestamp": time.time()
            }
        )


# Multi-Provider AI System Status Endpoint
@app.get("/api/providers/status")
@with_error_handling(ErrorCategory.SYSTEM, ErrorSeverity.LOW)
async def get_providers_status():
    """
    🤖 Multi-Provider AI System Status
    Shows the status of all configured AI providers and their role mappings.
    """
    try:
        factory = get_factory()

        # Get provider statistics
        stats = factory.get_provider_stats()

        # Get health status
        health_status = await factory.health_check_all()

        return {
            "status": "operational",
            "timestamp": time.time(),
            "multi_provider_system": {
                "total_providers": stats["total_providers"],
                "available_providers": stats["available_providers"],
                "provider_health": health_status,
                "role_mappings": stats["role_mappings"],
                "provider_models": stats["provider_models"]
            },
            "capabilities": {
                "dynamic_routing": True,
                "role_based_selection": True,
                "fallback_support": True,
                "streaming_support": True,
                "multi_model_support": True
            },
            "usage_examples": {
                "specific_provider": "/api/chat/provider/gemini",
                "role_based": "/api/chat/role/judge",
                "auto_routing": "/api/chat (default behavior)"
            }
        }

    except Exception as e:
        logger.error(f"Provider status check failed: {e}")
        return {
            "status": "error",
            "timestamp": time.time(),
            "error": str(e),
            "multi_provider_system": {
                "total_providers": 0,
                "available_providers": [],
                "provider_health": {}
            }
        }


# Provider-specific chat endpoint
@app.post("/api/chat/provider/{provider_name}")
@require_rate_limit("chat_request")
@validate_input()
@with_error_handling(ErrorCategory.AGENT_FAILURE, ErrorSeverity.MEDIUM)
async def chat_with_specific_provider(
    provider_name: str,
    message: str = Form(...),
    model: Optional[str] = Form(None),
    temperature: float = Form(0.7),
    max_tokens: Optional[int] = Form(None),
    current_user: Optional[Dict[str, Any]] = Depends(get_current_user)
):
    """
    🎯 Chat with Specific AI Provider

    Allows direct communication with a specific AI provider (gemini, openai, anthropic).
    Useful for testing provider capabilities or when you need a specific model.
    """
    try:
        from src.models.base import ChatMessage

        # Get the specific provider
        provider = get_provider(provider_name=provider_name)

        # Create message
        messages = [ChatMessage(role="user", content=message)]

        # Get response
        response = await provider.chat(
            messages=messages,
            model=model,
            temperature=temperature,
            max_tokens=max_tokens
        )

        return {
            "success": True,
            "provider": response.provider,
            "model": response.model,
            "response": response.content,
            "usage": response.usage,
            "metadata": response.metadata
        }

    except Exception as e:
        logger.error(f"Provider-specific chat error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# Role-based chat endpoint
@app.post("/api/chat/role/{role}")
@require_rate_limit("chat_request")
@validate_input()
@with_error_handling(ErrorCategory.AGENT_FAILURE, ErrorSeverity.MEDIUM)
async def chat_with_role_based_provider(
    role: str,
    message: str = Form(...),
    temperature: float = Form(0.7),
    max_tokens: Optional[int] = Form(None),
    current_user: Optional[Dict[str, Any]] = Depends(get_current_user)
):
    """
    🎭 Chat with Role-Based Provider Selection

    Automatically selects the best AI provider for a specific role:
    - judge: Best reasoning for evaluation
    - lawyer: Complex legal reasoning
    - researcher: Fast research capabilities
    - writer: Best for long-form writing
    - reviewer: Detailed analysis
    - summarizer: Fast summarization
    - general: Balanced performance
    """
    try:
        from src.models.base import ChatMessage, ModelRole

        # Convert string role to ModelRole enum
        try:
            model_role = ModelRole(role.lower())
        except ValueError:
            available_roles = [r.value for r in ModelRole]
            raise HTTPException(
                status_code=400,
                detail=f"Invalid role '{role}'. Available roles: {available_roles}"
            )

        # Get provider for this role
        provider = get_provider(role=model_role)

        # Create message with role-specific system prompt
        role_prompts = {
            ModelRole.JUDGE: "You are an expert judge evaluating arguments and evidence with fairness and precision.",
            ModelRole.LAWYER: "You are an experienced lawyer providing legal analysis and reasoning.",
            ModelRole.RESEARCHER: "You are a thorough researcher gathering and analyzing information.",
            ModelRole.WRITER: "You are a skilled writer creating clear, engaging, and well-structured content.",
            ModelRole.REVIEWER: "You are a detailed reviewer providing comprehensive analysis and feedback.",
            ModelRole.SUMMARIZER: "You are an expert at creating concise, accurate summaries.",
            ModelRole.GENERAL: "You are a helpful AI assistant."
        }

        messages = [
            ChatMessage(role="system", content=role_prompts.get(model_role, "")),
            ChatMessage(role="user", content=message)
        ]

        # Get response using role-appropriate model
        model = provider.get_default_model(model_role)
        response = await provider.chat(
            messages=messages,
            model=model,
            temperature=temperature,
            max_tokens=max_tokens
        )

        return {
            "success": True,
            "role": role,
            "provider": response.provider,
            "model": response.model,
            "response": response.content,
            "usage": response.usage,
            "role_optimization": f"Selected {response.provider} with {response.model} for {role} role"
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Role-based chat error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# Routing analysis endpoint for development/debugging
@app.post("/api/analyze")
@with_error_handling(ErrorCategory.SYSTEM, ErrorSeverity.LOW)
async def analyze_request_complexity(
    message: str = Form(...),
    files: Optional[List[UploadFile]] = File(None),
    user_params: Optional[str] = Form(None)
):
    """
    🔍 Analyze Request Complexity - Development endpoint
    Shows how the routing system would handle a request without actually processing it.
    """
    try:
        # Process files for analysis
        processed_files = []
        if files:
            for file in files:
                content = await file.read()
                processed_files.append({
                    "filename": file.filename,
                    "size": len(content),
                    "type": file.content_type
                })

        # Parse user parameters
        parsed_user_params = {}
        if user_params:
            try:
                parsed_user_params = json.loads(user_params)
            except json.JSONDecodeError:
                pass

        # Analyze routing
        routing = await unified_processor.router.analyze_request(
            message, processed_files, parsed_user_params
        )

        # Get detailed analysis
        return {
            "message": message,
            "analysis": {
                "word_count": len(message.split()),
                "file_count": len(processed_files),
                "has_user_params": bool(parsed_user_params),
                "academic_indicators": [
                    keyword for keyword in [
                        "essay", "research", "academic", "citation", "thesis",
                        "dissertation", "literature review", "methodology"
                    ] if keyword.lower() in message.lower()
                ]
            },
            "routing_decision": routing,
            "explanation": {
                "simple": "Fast responses for quick questions and simple tasks",
                "advanced": "Full academic workflow with research, writing, and quality assurance",
                "hybrid": "Parallel processing with both systems for comprehensive results"
            },
            "estimated_processing_time": {
                "simple": "1-3 seconds",
                "advanced": "30-300 seconds",
                "hybrid": "30-300 seconds (parallel)"
            }
        }

    except Exception as e:
        logger.error(f"Request analysis failed: {e}")
        return {
            "error": str(e),
            "message": message,
            "routing_decision": {
                "system": "advanced",
                "complexity": 5.0,
                "reason": "analysis_failed",
                "confidence": 0.0
            }
        }


@app.get("/health/detailed")
@with_error_handling(ErrorCategory.SYSTEM, ErrorSeverity.MEDIUM)
async def detailed_health_check():
    """Detailed health check with system status."""
    health_data = {
        "status": "healthy",
        "timestamp": time.time(),
        "version": "2.0.0",
        "uptime_seconds": (time.time() - app.state.start_time) if hasattr(app.state, 'start_time') else 0,
        "services": {}
    }

    # Check Redis
    try:
        await redis_client.ping()
        health_data["services"]["redis"] = {"status": "healthy", "response_time": "< 1ms"}
    except Exception as e:
        health_data["services"]["redis"] = {"status": "unhealthy", "error": str(e)}
        health_data["status"] = "degraded"

    # Check Database
    try:
        if db_manager.health_check():
            health_data["services"]["database"] = {"status": "healthy", "response_time": "< 10ms"}
        else:
            health_data["services"]["database"] = {"status": "unhealthy", "error": "Health check failed"}
            health_data["status"] = "degraded"
    except Exception as e:
        health_data["services"]["database"] = {"status": "unhealthy", "error": str(e)}
        health_data["status"] = "degraded"

    # Check Error Handler
    try:
        error_stats = await error_handler.get_error_statistics()
        health_data["services"]["error_handler"] = {
            "status": "healthy",
            "total_errors": error_stats["total_errors"],
            "critical_errors": error_stats["critical_errors"]
        }
    except Exception as e:
        health_data["services"]["error_handler"] = {"status": "unhealthy", "error": str(e)}

    return health_data


@app.get("/metrics")
@with_error_handling(ErrorCategory.SYSTEM, ErrorSeverity.LOW)
async def get_system_metrics():
    """Get comprehensive system metrics."""
    try:
        # Get error handler statistics
        error_stats = await error_handler.get_error_statistics()

        # Get middleware statistics
        middleware_instance = None
        for middleware in app.user_middleware:
            if hasattr(middleware, 'cls') and middleware.cls.__name__ == 'RevolutionaryErrorMiddleware':
                middleware_instance = middleware
                break

        middleware_stats = {}
        if middleware_instance and hasattr(middleware_instance, 'get_middleware_stats'):
            middleware_stats = await middleware_instance.get_middleware_stats()

        return {
            "timestamp": time.time(),
            "system": {
                "version": "2.0.0",
                "uptime": error_stats.get("uptime_seconds", 0)
            },
            "errors": error_stats,
            "middleware": middleware_stats,
            "services": {
                "redis": {"status": "healthy" if await _check_redis() else "unhealthy"},
                "database": {"status": "healthy" if db_manager.health_check() else "unhealthy"}
            }
        }
    except Exception as e:
        logger.error(f"Failed to get system metrics: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve system metrics")


async def _check_redis() -> bool:
    """Check Redis health."""
    try:
        await redis_client.ping()
        return True
    except:
        return False


# Authentication and Security endpoints
@app.post("/api/auth/login")
@require_rate_limit("authentication")
@validate_input()
@with_error_handling(ErrorCategory.AUTHENTICATION, ErrorSeverity.MEDIUM)
async def login(
    request: Request,
    login_data: Dict[str, Any],
    user_repo=Depends(get_user_repository)
):
    """Authenticate user and return JWT token."""
    wallet_address = login_data.get("wallet_address")

    if not wallet_address:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Wallet address is required"
        )

    # Get or create user
    user = user_repo.get_user_by_wallet(wallet_address)
    if not user:
        # Create new user for first-time login
        user = user_repo.create_user(
            wallet_address=wallet_address,
            user_type="student",
            subscription_tier="free"
        )

    # Generate JWT token
    token = await security_service.generate_jwt_token(user.to_dict())

    logger.info(f"User authenticated: {wallet_address}")

    return {
        "access_token": token,
        "token_type": "bearer",
        "user": {
            "id": str(user.id),
            "wallet_address": user.wallet_address,
            "user_type": user.user_type.value if user.user_type else "student",
            "subscription_tier": user.subscription_tier,
            "credits_remaining": user.credits_remaining
        }
    }


@app.post("/api/auth/verify")
@with_error_handling(ErrorCategory.AUTHENTICATION, ErrorSeverity.LOW)
async def verify_token(current_user: Dict[str, Any] = Depends(get_current_user)):
    """Verify JWT token and return user information."""
    return {
        "valid": True,
        "user": current_user
    }


from src.services.chunking_service import get_chunking_service
from src.services.vector_storage import get_vector_storage
from src.services.embedding_service import get_embedding_service

# Unified Chat Endpoint with Intelligent Routing
from src.api.schemas.chat import ChatRequest, ChatResponse

@app.post("/api/chat/advanced", response_model=ChatResponse, status_code=202)
@require_rate_limit("chat_request")
@validate_input()
@with_error_handling(ErrorCategory.AGENT_FAILURE, ErrorSeverity.MEDIUM)
async def unified_chat_endpoint(
    req: ChatRequest,
    current_user: Optional[Dict[str, Any]] = Depends(get_current_user),
    chunking_service = Depends(get_chunking_service),
    embedding_service = Depends(get_embedding_service),
    vector_storage = Depends(get_vector_storage)
):
    """
    🚀 Unified Chat Endpoint with Intelligent Routing

    Automatically routes between:
    - Simple Gemini system (quick responses)
    - Advanced HandyWriterz system (academic writing)
    - Hybrid mode (both systems in parallel)

    Routing is based on request complexity analysis.
    """
    try:
        # Generate trace_id early for SSE events
        trace_id = str(uuid.uuid4())

        # Process uploaded files - REAL IMPLEMENTATION
        processed_files = []
        file_context = ""

        if req.file_ids:
            logger.info(f"Loading content for {len(req.file_ids)} files: {req.file_ids}")

            # Send file processing start event via SSEService
            try:
                from src.services.sse_service import get_sse_service
                await get_sse_service().publish_file_processing(
                    trace_id,
                    status="processing_files",
                    extra={
                        "file_count": len(req.file_ids),
                        "message": f"Processing {len(req.file_ids)} uploaded files..."
                    }
                )
            except Exception as sse_error:
                logger.warning(f"Failed to send file processing SSE event: {sse_error}")

            # Import and use the file content service
            from src.services.file_content_service import get_file_content_service
            file_service = get_file_content_service()

            try:
                # Load actual file contents
                file_contents = await file_service.load_file_contents(req.file_ids)
                logger.info(f"Successfully loaded {len(file_contents)} files")

                # Send file processing completion event via SSEService
                try:
                    successful_files = [fc for fc in file_contents if not fc.error]
                    failed_files = [fc for fc in file_contents if fc.error]

                    from src.services.sse_service import get_sse_service
                    await get_sse_service().publish_file_processing(
                        trace_id,
                        status="files_processed",
                        extra={
                            "successful_count": len(successful_files),
                            "failed_count": len(failed_files),
                            "message": f"Processed {len(successful_files)} files successfully" + (f", {len(failed_files)} failed" if failed_files else "")
                        }
                    )
                except Exception as sse_error:
                    logger.warning(f"Failed to send file completion SSE event: {sse_error}")

                # Format for prompt inclusion
                file_context = file_service.format_files_for_prompt(file_contents)

                # Store for unified processor
                processed_files = [
                    {
                        "file_id": fc.file_id,
                        "filename": fc.filename,
                        "content": fc.content,
                        "mime_type": fc.mime_type,
                        "size": fc.size,
                        "error": fc.error
                    }
                    for fc in file_contents
                ]

            except Exception as file_error:
                logger.error(f"Failed to load file contents: {file_error}")

                # Send file processing error event via SSEService
                try:
                    from src.services.sse_service import get_sse_service
                    await get_sse_service().publish_file_processing(
                        trace_id,
                        status="file_processing_error",
                        extra={
                            "error": str(file_error),
                            "message": f"Failed to process files: {file_error}"
                        }
                    )
                except Exception as sse_error:
                    logger.warning(f"Failed to send file error SSE event: {sse_error}")

                # Include error in context so user knows what happened
                file_context = f"\n=== FILE LOADING ERROR ===\nFailed to load uploaded files: {file_error}\n=== END ERROR ===\n"

        # Get user ID if available
        user_id = str(current_user.get("id")) if current_user else None

        # Optionally normalize user_params before processing (feature-gated)
        _settings = get_settings()
        normalized_params = req.user_params
        if getattr(_settings, "feature_params_normalization", False):
            try:
                from src.agent.routing.normalization import normalize_user_params, validate_user_params
                normalized_params = normalize_user_params(normalized_params or {})
                validate_user_params(normalized_params)
            except Exception:
                # do-not-harm: fall back silently
                normalized_params = req.user_params

        # Combine user prompt with file context
        enhanced_message = req.prompt
        if file_context:
            enhanced_message = f"{file_context}\n\nUSER REQUEST:\n{req.prompt}"
            logger.info(f"Enhanced message with {len(processed_files)} files, total length: {len(enhanced_message)}")

        # Process using unified processor with streaming support
        result = await unified_processor.process_message(
            message=enhanced_message,
            files=processed_files,
            user_params=normalized_params,
            user_id=user_id,
            conversation_id=trace_id
        )

        # Use the trace_id we generated (conversation_id should match)

        # Enhanced response format compatible with both frontend expectations
        response = {
            "success": result.get("success", True),
            "trace_id": trace_id,  # Always provide trace_id for frontend
            "response": result.get("response", ""),
            "sources": result.get("sources", []),
            "workflow_status": result.get("workflow_status", "completed"),

            # Routing information
            "system_used": result.get("system_used", "unknown"),
            "complexity_score": result.get("complexity_score", 0.0),
            "routing_reason": result.get("routing_reason", ""),
            "routing_confidence": result.get("routing_confidence", 0.0),
            "processing_time": result.get("processing_time", 0.0),

            # Advanced features (when available)
            "conversation_id": trace_id,  # Use trace_id consistently
            "quality_score": result.get("quality_score"),
            "agent_metrics": result.get("agent_metrics", {}),
            "citation_count": result.get("citation_count", 0),
            "system_type": result.get("system_type", ""),

            # Hybrid mode specific
            "simple_insights": result.get("simple_insights"),
            "research_depth": result.get("research_depth", 0),
            "hybrid_results": result.get("hybrid_results", {})
        }

        # Log successful routing
        logger.info("✅ Chat processed successfully:")
        logger.info(f"   System: {result.get('system_used')}")
        logger.info(f"   Complexity: {result.get('complexity_score'):.1f}")
        logger.info(f"   Time: {result.get('processing_time'):.2f}s")

        return response

    except Exception as e:
        logger.error(f"Unified chat endpoint error: {e}")

        # Enhanced error response
        error_response = {
            "success": False,
            "response": f"I encountered an error processing your request: {str(e)}",
            "sources": [],
            "workflow_status": "failed",
            "system_used": "error_fallback",
            "complexity_score": 0.0,
            "error_details": {
                "error_type": type(e).__name__,
                "error_message": str(e)
            }
        }

        # Return error response instead of raising HTTP exception for better UX
        return error_response


# Quick chat endpoint for simple queries (explicit routing)
@app.post("/api/chat/simple")
async def simple_chat_endpoint_removed():
    """Simple chat endpoint removed - use /api/chat which routes to advanced system."""
    raise HTTPException(
        status_code=410,
        detail="Simple chat system has been removed. Please use /api/chat endpoint which provides superior academic writing capabilities."
    )


# Advanced chat endpoint for academic writing (explicit routing)
@app.post("/api/chat/advanced")
@require_rate_limit("create_document")
@validate_input()
@with_error_handling(ErrorCategory.AGENT_FAILURE, ErrorSeverity.HIGH)
async def advanced_chat_endpoint(
    message: str = Form(...),
    files: Optional[List[UploadFile]] = File(None),
    user_params: Optional[str] = Form(None),
    current_user: Optional[Dict[str, Any]] = Depends(get_current_user)
):
    """
    Advanced chat endpoint - forces routing to HandyWriterz system for academic writing.
    Use this for complex academic writing tasks.
    """
    try:
        # Process files
        processed_files = []
        if files:
            for file in files:
                content = await file.read()
                processed_files.append({
                    "filename": file.filename,
                    "content": content.decode('utf-8', errors='ignore') if file.content_type and file.content_type.startswith("text") else content
                })

        # Parse user parameters
        parsed_user_params = {}
        if user_params:
            try:
                parsed_user_params = json.loads(user_params)
            except json.JSONDecodeError:
                logger.warning(f"Invalid user_params JSON: {user_params}")

        user_id = str(current_user.get("id")) if current_user else None

        # Force advanced processing
        result = await unified_processor._process_advanced(
            message, processed_files, parsed_user_params, user_id
        )

        return {
            "success": True,
            "response": result.get("response", ""),
            "conversation_id": result.get("conversation_id"),
            "sources": result.get("sources", []),
            "quality_score": result.get("quality_score", 0),
            "system_used": "advanced_forced",
            "processing_time": 0.0  # Would need to time this
        }

    except Exception as e:
        logger.error(f"Advanced chat endpoint error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/security/stats")
@with_error_handling(ErrorCategory.SYSTEM, ErrorSeverity.LOW)
async def get_security_stats(
    current_user: Dict[str, Any] = Depends(require_authorization("admin_access"))
):
    """Get security statistics (admin only)."""
    try:
        # Get security events from Redis
        security_events = await redis_client.lrange("security_events", 0, 99)

        # Get middleware stats
        middleware_stats = {}
        for middleware in app.user_middleware:
            if hasattr(middleware, 'cls') and hasattr(middleware.cls, 'get_security_stats'):
                instance = getattr(middleware, 'instance', None)
                if instance and hasattr(instance, 'get_security_stats'):
                    middleware_stats = await instance.get_security_stats()
                    break

        return {
            "security_events": [json.loads(event) for event in security_events],
            "middleware_stats": middleware_stats,
            "total_events": len(security_events),
            "timestamp": time.time()
        }

    except Exception as e:
        logger.error(f"Failed to get security stats: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to retrieve security statistics"
        )


# Main writing endpoint
@app.post("/api/write", response_model=WritingResponse)
@require_rate_limit("create_document")
@validate_input()
@with_circuit_breaker("writing_workflow")
@with_error_handling(ErrorCategory.AGENT_FAILURE, ErrorSeverity.HIGH)
async def start_writing(
    request: WritingRequest,
    http_request: Request,
    current_user: Optional[Dict[str, Any]] = Depends(get_current_user),
    user_repo=Depends(get_user_repository),
    conversation_repo=Depends(get_conversation_repository)
):
    """Start the academic writing process with comprehensive error handling."""
    context = ErrorContext(
        request_id=getattr(http_request.state, 'request_id', str(uuid.uuid4())),
        additional_data={
            "prompt_length": len(request.prompt),
            "user_params": request.user_params,
            "file_count": len(request.uploaded_file_urls)
        }
    )

    try:
        # Create or get user from wallet address
        user = None
        if request.auth_token:  # If user is authenticated
            # Extract wallet address from auth token (simplified)
            wallet_address = request.auth_token  # Placeholder - implement proper JWT parsing
            user = user_repo.get_user_by_wallet(wallet_address)
            if not user:
                user = user_repo.create_user(
                    wallet_address=wallet_address,
                    user_type="student",
                    subscription_tier="free"
                )

        # Optionally normalize user parameters (feature-gated) before Pydantic validation
        _settings = get_settings()
        incoming_params = request.user_params or {}

        # Check feature flag from settings or environment
        feature_enabled = (
            getattr(_settings, "feature_params_normalization", False) if _settings
            else os.getenv("FEATURE_PARAMS_NORMALIZATION", "false").lower() == "true"
        )

        if feature_enabled:
            try:
                from src.agent.routing.normalization import normalize_user_params, validate_user_params
                logger.debug(f"Normalizing user params: {incoming_params}")
                normalized_params = normalize_user_params(incoming_params)
                validate_user_params(normalized_params)
                incoming_params = normalized_params
                logger.debug(f"Successfully normalized params: {normalized_params}")
            except Exception as e:
                # Keep original on failure to avoid harm
                logger.warning(f"Parameter normalization failed, using original: {e}")
                incoming_params = request.user_params or {}

        # Validate user parameters
        try:
            user_params = UserParams(**incoming_params)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Invalid user parameters: {e}")

        # Create conversation in database
        conversation = conversation_repo.create_conversation(
            user_id=str(user.id) if user else None,
            user_params=user_params.dict(),
            title=f"{user_params.writeupType.title()} - {user_params.field.title()}",
            workflow_status="initiated"
        )

        conversation_id = str(conversation.id)

        # Create initial state
        initial_state = HandyWriterzState(
            conversation_id=conversation_id,
            user_id=str(user.id) if user else "",
            wallet_address=user.wallet_address if user else None,
            messages=[HumanMessage(content=request.prompt)],
            user_params=user_params.dict(),
            uploaded_docs=[],
            outline=None,
            research_agenda=[],
            search_queries=[],
            raw_search_results=[],
            filtered_sources=[],
            verified_sources=[],
            draft_content=None,
            current_draft=None,
            revision_count=0,
            evaluation_results=[],
            evaluation_score=None,
            turnitin_reports=[],
            turnitin_passed=False,
            formatted_document=None,
            learning_outcomes_report=None,
            download_urls={},
            current_node=None,
            workflow_status="initiated",
            error_message=None,
            retry_count=0,
            max_iterations=5,
            enable_tutor_review=False,
            start_time=time.time(),
            end_time=None,
            processing_metrics={},
            auth_token=request.auth_token,
            payment_transaction_id=request.payment_transaction_id,
            uploaded_files=[{"url": url} for url in request.uploaded_file_urls]
        )

        # Start the workflow asynchronously
        asyncio.create_task(execute_writing_workflow(conversation_id, initial_state))

        return WritingResponse(
            conversation_id=conversation_id,
            status="started",
            message="Revolutionary academic writing process initiated. Connect to the stream endpoint for real-time updates."
        )

    except Exception as e:
        logger.error(f"Failed to start writing process: {e}")
        raise HTTPException(status_code=500, detail=str(e))


## Legacy SSE endpoint removed in favor of routes/stream.py canonical SSE router
## (kept here as comment to avoid accidental reintroduction)
#                        # Some publishers may already send dicts; ensure string -> dict
#                        event_data: Dict[str, Any]
#                        if isinstance(data_raw, str):
#                            event_data = json.loads(data_raw)
#                        elif isinstance(data_raw, dict):
#                            event_data = data_raw
#                        else:
#                            # Unknown payload type; skip
#                            continue
#
#                        # Ensure timestamp for client ordering if missing
#                        if "ts" not in event_data:
#                            event_data["ts"] = time.time()
#
#                        # Normalize legacy 'text' to 'content'
#                        if "text" in event_data and "content" not in event_data:
#                            event_data["content"] = event_data.pop("text")
#
#                        yield f"data: {json.dumps(event_data)}\n\n"
#
#                        # Close on terminal events
## Legacy SSE endpoint removed in favor of routes/stream.py canonical SSE router
## (kept here as comment to avoid accidental reintroduction)
#                        # Some publishers may already send dicts; ensure string -> dict
#                        event_data: Dict[str, Any]
#                        if isinstance(data_raw, str):
#                            event_data = json.loads(data_raw)
#                        elif isinstance(data_raw, dict):
#                            event_data = data_raw
#                        else:
#                            # Unknown payload type; skip
#                            continue
#
#                        # Ensure timestamp for client ordering if missing
#                        if "ts" not in event_data:
#                            event_data["ts"] = time.time()
#
#                        # Normalize legacy 'text' to 'content'
#                        if "text" in event_data and "content" not in event_data:
#                            event_data["content"] = event_data.pop("text")
#
#                        yield f"data: {json.dumps(event_data)}\n\n"
#
#                        # Close on terminal events
#                        evt_type = event_data.get("type")
#                        if evt_type in ["workflow_complete", "workflow_failed", "workflow_finished", "done", "error"]:
#                            break
#
#                    except Exception as e:
#                        logger.error(f"Error processing SSE message: {e}")
#                        # Send an error event but keep the stream alive
#                        yield f"data: {json.dumps({'type': 'error', 'message': f'Event processing error: {str(e)}', 'ts': time.time()})}\n\n"
#                        continue
#
#                except Exception as e:
#                    logger.error(f"SSE stream error: {e}")
#                    yield f"data: {json.dumps({'type': 'error', 'message': str(e), 'ts': time.time()})}\n\n"
#
#                finally:
#                    # Always attempt to unsubscribe and close pubsub
#                    try:
#                        await pubsub.unsubscribe(channel)  # type: ignore[reportUnknownMemberType]
#                    except Exception:
#                        pass
#                    try:
#                        await pubsub.close()  # type: ignore[reportUnknownMemberType]
#                    except Exception:
#                        pass
#                    logger.info(f"Unsubscribed from SSE channel: {channel} (disconnected={disconnected})")

# Strong anti-buffering and keep-alive headers
#headers = {
#    "Cache-Control": "no-cache, no-transform",
#    "Pragma": "no-cache",
#    "Connection": "keep-alive",
#    "X-Accel-Buffering": "no",
#    "Keep-Alive": "timeout=60, max=1000",
#    # Content-Type is set by media_type below
#    # Transfer-Encoding is set automatically by ASGI server for streaming
#}
#
#return StreamingResponse(
#    generate_events(),
#    media_type="text/event-stream; charset=utf-8",
#    headers=headers
#)




@app.post("/api/retrieve")
async def retrieve(
    query: str,
    k: int = 10,
    embedding_service = Depends(get_embedding_service),
    vector_storage = Depends(get_vector_storage)
):
    query_embedding = await embedding_service.embed_query(query)
    results = await vector_storage.retrieve_chunks(query_embedding, k)
    return {"results": results}


# Get conversation status
@app.get("/api/conversation/{conversation_id}")
async def get_conversation_status(
    conversation_id: str,
    conversation_repo=Depends(get_conversation_repository)
):
    """Get the current status of a conversation."""
    try:
        # Query database for real conversation status
        conversation = conversation_repo.get_conversation(conversation_id)

        if not conversation:
            raise HTTPException(
                status_code=404,
                detail=f"Conversation {conversation_id} not found"
            )

        # Calculate progress based on workflow status
        progress_map = {
            "initiated": 5.0,
            "planning": 15.0,
            "searching": 35.0,
            "filtering": 50.0,
            "writing": 70.0,
            "evaluating": 85.0,
            "formatting": 95.0,
            "completed": 100.0,
            "failed": 0.0
        }

        current_progress = progress_map.get(conversation.workflow_status, 0.0)

        # Estimate completion time based on status
        time_estimates = {
            "initiated": "8-12 minutes",
            "planning": "6-10 minutes",
            "searching": "4-8 minutes",
            "filtering": "3-5 minutes",
            "writing": "2-4 minutes",
            "evaluating": "1-2 minutes",
            "formatting": "30 seconds",
            "completed": "Complete",
            "failed": "Failed"
        }

        return {
            "conversation_id": conversation_id,
            "status": conversation.workflow_status,
            "current_node": conversation.current_node,
            "progress": current_progress,
            "estimated_completion": time_estimates.get(conversation.workflow_status, "Unknown"),
            "created_at": conversation.created_at.isoformat(),
            "updated_at": conversation.updated_at.isoformat(),
            "user_params": conversation.user_params,
            "error_message": conversation.error_message,
            "retry_count": conversation.retry_count or 0
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get conversation status: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# Download document endpoint
@app.get("/api/download/{conversation_id}/{document_type}")
async def download_document(
    conversation_id: str,
    document_type: str,
    document_repo=Depends(get_document_repository)
):
    """Download a generated document."""
    try:
        # Validate document type
        allowed_types = ["docx", "txt", "pdf", "lo_report"]
        if document_type not in allowed_types:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid document type: {document_type}. Allowed: {allowed_types}"
            )

        # Get document from database
        document = document_repo.get_by_conversation_and_type(conversation_id, document_type)

        if not document:
            raise HTTPException(
                status_code=404,
                detail=f"Document of type {document_type} not found for conversation {conversation_id}"
            )

        # Check if document has a file URL (cloud storage)
        if document.file_urls and document_type in document.file_urls:
            file_url = document.file_urls[document_type]

            # For cloud storage, redirect to the file URL
            from fastapi.responses import RedirectResponse
            return RedirectResponse(url=file_url, status_code=302)

        # If no file URL, generate document content on-the-fly
        content = document.content
        if not content:
            raise HTTPException(
                status_code=404,
                detail="Document content not available"
            )

        # Generate appropriate response based on document type
        if document_type == "docx":
            from io import BytesIO
            from docx import Document as DocxDocument

            # Create DOCX from content
            doc = DocxDocument()
            for paragraph in content.split('\n\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

            bio = BytesIO()
            doc.save(bio)
            bio.seek(0)

            from fastapi.responses import StreamingResponse
            return StreamingResponse(
                BytesIO(bio.getvalue()),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={conversation_id}_{document_type}.docx"}
            )

        elif document_type == "txt":
            from fastapi.responses import Response
            return Response(
                content=content,
                media_type="text/plain",
                headers={"Content-Disposition": f"attachment; filename={conversation_id}_{document_type}.txt"}
            )

        elif document_type == "lo_report":
            # Learning outcomes report as JSON
            lo_data = document.metadata.get("learning_outcomes", {}) if document.metadata else {}
            import json

            from fastapi.responses import Response
            return Response(
                content=json.dumps(lo_data, indent=2),
                media_type="application/json",
                headers={"Content-Disposition": f"attachment; filename={conversation_id}_learning_outcomes.json"}
            )

        else:
            raise HTTPException(status_code=400, detail=f"Unsupported document type: {document_type}")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document download failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# User management endpoints
@app.get("/api/users/{wallet_address}")
async def get_user_profile(
    wallet_address: str,
    user_repo=Depends(get_user_repository)
):
    """Get user profile by wallet address."""
    try:
        user = user_repo.get_user_by_wallet(wallet_address)

        if not user:
            raise HTTPException(
                status_code=404,
                detail="User not found"
            )

        return {
            "id": str(user.id),
            "wallet_address": user.wallet_address,
            "user_type": user.user_type,
            "subscription_tier": user.subscription_tier,
            "credits_balance": user.credits_balance,
            "credits_used": user.credits_used,
            "documents_created": user.documents_created,
            "avg_quality_score": user.avg_quality_score,
            "created_at": user.created_at.isoformat(),
            "last_active": user.last_active.isoformat() if user.last_active else None,
            "preferences": user.preferences
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get user profile: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/api/users/{wallet_address}")
async def update_user_profile(
    wallet_address: str,
    user_data: Dict[str, Any],
    user_repo=Depends(get_user_repository)
):
    """Update user profile."""
    try:
        user = user_repo.get_user_by_wallet(wallet_address)

        if not user:
            raise HTTPException(
                status_code=404,
                detail="User not found"
            )

        # Update allowed fields
        allowed_fields = ['user_type', 'subscription_tier', 'preferences']
        update_data = {k: v for k, v in user_data.items() if k in allowed_fields}

        user_repo.update_user_stats(str(user.id), **update_data)

        return {"status": "updated", "message": "User profile updated successfully"}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to update user profile: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/users/{wallet_address}/conversations")
async def get_user_conversations(
    wallet_address: str,
    page: int = 1,
    page_size: int = 20,
    sort_by: str = "created_at",
    sort_dir: str = "desc",
    user_repo=Depends(get_user_repository),
    conversation_repo=Depends(get_conversation_repository)
):
    """Get user's conversation history with pagination and sorting."""
    try:
        user = user_repo.get_user_by_wallet(wallet_address)

        if not user:
            raise HTTPException(
                status_code=404,
                detail="User not found"
            )

        # Sanitize page params
        page = max(1, page)
        page_size = max(1, min(100, page_size))
        offset = (page - 1) * page_size

        total = conversation_repo.count_user_conversations(str(user.id))
        conversations = conversation_repo.get_user_conversations(
            str(user.id), limit=page_size, offset=offset, sort_by=sort_by, sort_dir=sort_dir
        )

        return {
            "data": [
                {
                    "id": str(conv.id),
                    "title": conv.title,
                    "workflow_status": conv.workflow_status,
                    "created_at": conv.created_at.isoformat(),
                    "updated_at": conv.updated_at.isoformat(),
                    "user_params": conv.user_params,
                    "error_message": conv.error_message
                }
                for conv in conversations
            ],
            "meta": {
                "page": page,
                "page_size": page_size,
                "total": total,
                "total_pages": (total + page_size - 1) // page_size,
                "sort_by": sort_by,
                "sort_dir": sort_dir
            }
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get user conversations: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# Credits management endpoints
@app.get("/api/credits/{wallet_address}")
async def get_user_credits(
    wallet_address: str,
    user_repo=Depends(get_user_repository)
):
    """Get user's credit balance and usage statistics."""
    try:
        user = user_repo.get_user_by_wallet(wallet_address)

        if not user:
            # Create user if they don't exist
            user = user_repo.create_user(
                wallet_address=wallet_address,
                user_type="student",
                subscription_tier="free",
                credits_balance=3  # Welcome bonus
            )

        return {
            "wallet_address": wallet_address,
            "credits_balance": user.credits_balance,
            "credits_used": user.credits_used,
            "subscription_tier": user.subscription_tier,
            "next_renewal": None,  # TODO: Add subscription logic
            "usage_stats": {
                "documents_created": user.documents_created,
                "avg_quality_score": user.avg_quality_score,
                "total_words_generated": user.total_words_generated or 0
            }
        }

    except Exception as e:
        logger.error(f"Failed to get user credits: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/credits/{wallet_address}/purchase")
async def purchase_credits(
    wallet_address: str,
    purchase_data: Dict[str, Any],
    user_repo=Depends(get_user_repository)
):
    """Purchase credits for a user."""
    try:
        amount = purchase_data.get("amount", 0)
        payment_method = purchase_data.get("payment_method", "crypto")
        transaction_id = purchase_data.get("transaction_id")

        if amount <= 0:
            raise HTTPException(
                status_code=400,
                detail="Invalid credit amount"
            )

        user = user_repo.get_user_by_wallet(wallet_address)

        if not user:
            raise HTTPException(
                status_code=404,
                detail="User not found"
            )

        # TODO: Verify payment transaction
        # For now, we'll trust the transaction
        # Add credits to user balance
        new_balance = user.credits_balance + amount
        user_repo.update_user_stats(
            str(user.id),
            credits_balance=new_balance
        )

        logger.info(f"Credits purchased: {wallet_address} - {amount} credits")

        return {
            "status": "success",
            "credits_added": amount,
            "new_balance": new_balance,
            "transaction_id": transaction_id
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to purchase credits: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/credits/{wallet_address}/bonus")
async def claim_welcome_bonus(
    wallet_address: str,
    user_repo=Depends(get_user_repository)
):
    """Claim welcome bonus credits."""
    try:
        user = user_repo.get_user_by_wallet(wallet_address)

        if not user:
            # Create user with welcome bonus
            user = user_repo.create_user(
                wallet_address=wallet_address,
                user_type="student",
                subscription_tier="free",
                credits_balance=3,  # Welcome bonus
                welcome_bonus_claimed=True
            )

            return {
                "status": "success",
                "credits_granted": 3,
                "new_balance": 3,
                "message": "Welcome bonus claimed!"
            }

        # Check if bonus already claimed
        if user.welcome_bonus_claimed:
            raise HTTPException(
                status_code=400,
                detail="Welcome bonus already claimed"
            )

        # Grant welcome bonus
        new_balance = user.credits_balance + 3
        user_repo.update_user_stats(
            str(user.id),
            credits_balance=new_balance,
            welcome_bonus_claimed=True
        )

        logger.info(f"Welcome bonus claimed: {wallet_address}")

        return {
            "status": "success",
            "credits_granted": 3,
            "new_balance": new_balance,
            "message": "Welcome bonus claimed!"
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to claim welcome bonus: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# Webhook endpoints
@app.post("/api/webhook/dynamic")
async def dynamic_webhook(payload: Dict[str, Any]):
    """Handle Dynamic.xyz webhooks for payment verification."""
    try:
        # Process Dynamic.xyz webhook
        event_type = payload.get("type")

        if event_type == "payment.completed":
            # Handle successful payment
            logger.info(f"Payment completed: {payload}")

        elif event_type == "user.authenticated":
            # Handle user authentication
            logger.info(f"User authenticated: {payload}")

        return {"status": "received"}

    except Exception as e:
        logger.error(f"Dynamic webhook processing failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/webhook/turnitin")
async def turnitin_webhook(payload: Dict[str, Any]):
    """Handle Turnitin webhooks for plagiarism reports."""
    try:
        # Process Turnitin webhook
        submission_id = payload.get("submission_id")
        status = payload.get("status")

        if status == "completed":
            # Handle completed plagiarism check
            logger.info(f"Turnitin check completed for submission: {submission_id}")

        return {"status": "received"}

    except Exception as e:
        logger.error(f"Turnitin webhook processing failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# Lightweight readiness endpoint
@app.get("/health/ready")
async def health_ready():
    """Fast readiness probe with Redis and DB checks.

    Returns: {"redis":"ok|fail","db":"ok|fail","version":"<sha|dev>"}
    """
    # Redis
    try:
        await asyncio.wait_for(redis_client.ping(), timeout=1.5)  # keep <150ms target when healthy
        redis_status = "ok"
    except Exception:
        redis_status = "fail"

    # DB
    try:
        db_ok = db_manager.health_check()
        db_status = "ok" if db_ok else "fail"
    except Exception:
        db_status = "fail"

    # Version from env; fallback to dev
    version = os.getenv("GIT_SHA") or os.getenv("RELEASE_SHA") or "dev"

    return {"redis": redis_status, "db": db_status, "version": version}


# Import vector storage dependencies
from src.services.vector_storage import get_vector_storage
from src.services.embedding_service import get_embedding_service


# Vector search endpoints
@app.post("/api/search/semantic")
async def semantic_search(
    request: Dict[str, Any],
    vector_storage=Depends(get_vector_storage),
    embedding_service=Depends(get_embedding_service)
):
    """Perform semantic search using vector embeddings."""
    try:
        query = request.get("query", "")
        if not query:
            raise HTTPException(status_code=400, detail="Query is required")

        # Optional parameters
        academic_field = request.get("academic_field")
        min_credibility = request.get("min_credibility", 0.6)
        limit = min(request.get("limit", 10), 50)  # Max 50 results
        year_range = request.get("year_range")  # [start_year, end_year]

        # Generate query embedding
        query_embedding = await embedding_service.embed_query(query, "academic_search")

        # Perform semantic search
        results = await vector_storage.semantic_search_documents(
            query_embedding=query_embedding,
            limit=limit,
            academic_field=academic_field,
            min_credibility=min_credibility,
            year_range=tuple(year_range) if year_range and len(year_range) == 2 else None
        )

        return {
            "query": query,
            "results": results,
            "total_found": len(results),
            "search_metadata": {
                "academic_field": academic_field,
                "min_credibility": min_credibility,
                "year_range": year_range,
                "timestamp": time.time()
            }
        }

    except Exception as e:
        logger.error(f"Semantic search failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/evidence/{conversation_id}")
async def get_conversation_evidence(
    conversation_id: str,
    vector_storage=Depends(get_vector_storage)
):
    """Get all evidence for a conversation with vector similarity support."""
    try:
        evidence_list = await vector_storage.get_conversation_evidence(conversation_id)

        return {
            "conversation_id": conversation_id,
            "evidence_count": len(evidence_list),
            "evidence": evidence_list
        }

    except Exception as e:
        logger.error(f"Failed to get conversation evidence: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/download/{conversation_id}")
async def list_available_downloads(
    conversation_id: str,
    document_repo=Depends(get_document_repository)
):
    """List all available download formats for a conversation."""
    try:
        documents = document_repo.get_conversation_documents(conversation_id)

        if not documents:
            raise HTTPException(
                status_code=404,
                detail=f"No documents found for conversation {conversation_id}"
            )

        available_downloads = []

        for doc in documents:
            download_info = {
                "document_id": str(doc.id),
                "title": doc.title,
                "document_type": doc.document_type.value if doc.document_type else "unknown",
                "word_count": doc.word_count,
                "quality_score": doc.overall_quality_score,
                "available_formats": []
            }

            # Check available formats
            if doc.content_markdown:
                download_info["available_formats"].append({
                    "format": "txt",
                    "url": f"/api/download/{conversation_id}/txt",
                    "description": "Plain text format"
                })

            if doc.docx_url or doc.content_markdown:
                download_info["available_formats"].append({
                    "format": "docx",
                    "url": f"/api/download/{conversation_id}/docx",
                    "description": "Microsoft Word format"
                })

            if doc.pdf_url:
                download_info["available_formats"].append({
                    "format": "pdf",
                    "url": f"/api/download/{conversation_id}/pdf",
                    "description": "PDF format"
                })

            if doc.learning_outcomes_coverage:
                download_info["available_formats"].append({
                    "format": "lo_report",
                    "url": f"/api/download/{conversation_id}/lo_report",
                    "description": "Learning outcomes report (JSON)"
                })

            available_downloads.append(download_info)

        return {
            "conversation_id": conversation_id,
            "documents": available_downloads
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to list downloads: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/download/{conversation_id}/zip")
async def download_all_formats(
    conversation_id: str,
    document_repo=Depends(get_document_repository)
):
    """Download all available formats as a ZIP file."""
    try:
        import zipfile
        from io import BytesIO
        from docx import Document as DocxDocument

        documents = document_repo.get_conversation_documents(conversation_id)

        if not documents:
            raise HTTPException(
                status_code=404,
                detail=f"No documents found for conversation {conversation_id}"
            )

        # Create ZIP file in memory
        zip_buffer = BytesIO()

        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:

            for doc in documents:
                base_name = f"{doc.title.replace(' ', '_')[:50]}"

                # Add TXT format
                if doc.content_markdown:
                    zip_file.writestr(f"{base_name}.txt", doc.content_markdown)

                # Add DOCX format
                if doc.content_markdown:
                    docx_doc = DocxDocument()
                    for paragraph in doc.content_markdown.split('\n\n'):
                        if paragraph.strip():
                            docx_doc.add_paragraph(paragraph.strip())

                    docx_buffer = BytesIO()
                    docx_doc.save(docx_buffer)
                    zip_file.writestr(f"{base_name}.docx", docx_buffer.getvalue())

                # Add learning outcomes report
                if doc.learning_outcomes_coverage:
                    import json
                    lo_json = json.dumps(doc.learning_outcomes_coverage, indent=2)
                    zip_file.writestr(f"{base_name}_learning_outcomes.json", lo_json)

                # Add metadata
                metadata = {
                    "title": doc.title,
                    "word_count": doc.word_count,
                    "quality_score": doc.overall_quality_score,
                    "citation_count": doc.citation_count,
                    "academic_field": doc.academic_field,
                    "generated_at": doc.created_at.isoformat() if doc.created_at else None
                }
                zip_file.writestr(f"{base_name}_metadata.json", json.dumps(metadata, indent=2))

        zip_buffer.seek(0)

        from fastapi.responses import StreamingResponse
        return StreamingResponse(
            BytesIO(zip_buffer.getvalue()),
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename={conversation_id}_complete.zip"}
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"ZIP download failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# Background workflow execution with comprehensive error handling
@with_retry(ErrorCategory.AGENT_FAILURE)
async def execute_writing_workflow(conversation_id: str, initial_state: HandyWriterzState):
    """Execute the writing workflow with production-grade error handling."""
    context = ErrorContext(
        conversation_id=conversation_id,
        user_id=initial_state.user_id,
        additional_data={
            "workflow_type": "academic_writing",
            "user_params": initial_state.user_params
        }
    )

    # Import SSE service lazily to avoid circulars at import time
    from src.services.sse_service import get_sse_service
    sse = get_sse_service()

    try:
        logger.info(f"🚀 Starting revolutionary workflow for conversation: {conversation_id}")

        # Broadcast workflow start with enhanced data
        await sse.publish_workflow_start(conversation_id, {
            "conversation_id": conversation_id,
            "user_id": initial_state.user_id,
            "estimated_duration": "8-12 minutes",
            "workflow_version": "2.0.0"
        })

        config = {"configurable": {"thread_id": conversation_id}}

        workflow_start_time = time.time()
        chunk_count = 0

        # Import the graph lazily to avoid import-order issues during startup
        from src.agent.handywriterz_graph import handywriterz_graph

        # Execute the LangGraph workflow with circuit breaker
        async for chunk in handywriterz_graph.astream(initial_state, config):
            chunk_count += 1

            # Enhanced progress broadcasting
            await sse.publish_workflow_progress(conversation_id, {
                **chunk,
                "chunk_number": chunk_count,
                "elapsed_time": time.time() - workflow_start_time
            })

            # Log major progress milestones
            if "current_node" in chunk:
                logger.info(f"📍 Workflow [{conversation_id}] progressed to: {chunk['current_node']}")

        workflow_duration = time.time() - workflow_start_time

        # Broadcast successful completion
        await sse.publish_workflow_complete(conversation_id, {
            "conversation_id": conversation_id,
            "status": "completed",
            "duration_seconds": workflow_duration,
            "chunks_processed": chunk_count,
            "completion_message": "Academic document generated successfully."
        })

        logger.info(f"✅ Workflow completed successfully for {conversation_id} in {workflow_duration:.2f}s")

    except Exception as e:
        workflow_duration = time.time() - workflow_start_time if 'workflow_start_time' in locals() else 0

        # Handle error through error handler
        error_data = await error_handler.handle_error(
            e, context, ErrorCategory.AGENT_FAILURE, ErrorSeverity.HIGH
        )

        logger.error(f"❌ Workflow execution failed for {conversation_id}: {e}")

        # Determine if error is recoverable
        recovery_strategy = error_data.get("recovery_strategy", {})
        is_recoverable = recovery_strategy.get("retry_recommended", False)

        # Broadcast workflow failure with recovery information
        await sse.publish_workflow_failed(conversation_id, {
            "conversation_id": conversation_id,
            "error": str(e),
            "error_id": error_data.get("error_id"),
            "error_type": type(e).__name__,
            "duration_seconds": workflow_duration,
            "recoverable": is_recoverable,
            "recovery_strategy": recovery_strategy,
            "support_message": "Our team has been notified. Please try again or contact support."
        })

        # Re-raise if not recoverable
        if not is_recoverable:
            raise


# Configuration endpoint for frontend compatibility
@app.get("/api/config")
@with_error_handling(ErrorCategory.SYSTEM, ErrorSeverity.LOW)
async def get_app_config(request: Request):
    """
    Get application configuration for frontend.
    Compatible with OpenWebUI frontend expectations.
    """

    # For now, return a simplified config that satisfies the frontend
    # This can be expanded as needed
    return {
        "status": True,
        "name": "HandyWriterz",
        "version": "2.0.0",
        "default_locale": "en-US",
        "oauth": {
            "providers": {}
        },
        "features": {
            "auth": True,
            "auth_trusted_header": False,
            "enable_ldap": False,
            "enable_api_key": True,
            "enable_signup": True,
            "enable_login_form": True,
            "enable_websocket": True,
            "enable_direct_connections": False,
            "enable_channels": False,
            "enable_notes": True,
            "enable_web_search": True,
            "enable_code_execution": False,
            "enable_code_interpreter": False,
            "enable_image_generation": False,
            "enable_autocomplete_generation": True,
            "enable_community_sharing": False,
            "enable_message_rating": True,
            "enable_user_webhooks": False,
            "enable_admin_export": False,
            "enable_admin_chat_access": True,
            "enable_google_drive_integration": False,
            "enable_onedrive_integration": False,
        },
        "default_models": [],
        "default_prompt_suggestions": [
            {
                "title": "Academic Essay",
                "content": "Help me write a 2000-word academic essay on sustainable healthcare practices with Harvard referencing."
            },
            {
                "title": "Research Report",
                "content": "Create a comprehensive research report analyzing the impact of AI on modern education systems."
            },
            {
                "title": "Literature Review",
                "content": "Write a literature review examining recent developments in renewable energy technologies."
            }
        ],
        "user_count": 1,
        "code": {
            "engine": "none"
        },
        "audio": {
            "tts": {
                "engine": "none",
                "voice": "default",
                "split_on": "punctuation"
            },
            "stt": {
                "engine": "none"
            }
        },
        "file": {
            "max_size": 50,  # MB
            "max_count": 10,
            "image_compression": {
                "width": 1920,
                "height": 1080
            }
        },
        "permissions": {
            "workspace": {
                "models": True,
                "knowledge": True,
                "prompts": True,
                "tools": True,
                "functions": True
            },
            "chat": {
                "file_upload": True,
                "delete": True,
                "edit": True,
                "temporary": True
            }
        }
    }


# SPA Catch-all route to serve the frontend
@app.get("/")
async def serve_frontend():
    """Serve the SvelteKit build index.html file."""
    build_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "build")
    index_html = os.path.join(build_dir, "index.html")

    if os.path.exists(index_html):
        return FileResponse(index_html, media_type="text/html")
    else:
        # Fallback to development mode - serve from src
        frontend_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "src")
        app_html = os.path.join(frontend_dir, "app.html")

        if os.path.exists(app_html):
            return FileResponse(app_html, media_type="text/html")
        else:
            return JSONResponse(
                status_code=503,
                content={
                    "error": "Frontend not available",
                    "message": "Please build the frontend with 'npm run build' or run in development mode"
                }
            )

@app.get("/chat")
async def serve_chat():
    """Serve the chat page."""
    return await serve_frontend()

@app.get("/c/{chat_id}")
async def serve_chat_with_id(chat_id: str):
    """Serve specific chat page."""
    return await serve_frontend()

@app.get("/{path:path}")
async def serve_spa(path: str):
    """Catch-all route for SPA routing."""
    # Skip API routes
    if path.startswith("api/"):
        return JSONResponse(status_code=404, content={"error": "API endpoint not found"})

    # Skip static files
    if path.startswith("static/") or path.startswith("_app/") or path.startswith("pyodide/"):
        return JSONResponse(status_code=404, content={"error": "Static file not found"})

    # For all other routes, serve the SPA
    return await serve_frontend()


if __name__ == "__main__":
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="info"
    )
